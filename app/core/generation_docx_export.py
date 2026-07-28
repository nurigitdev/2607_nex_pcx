"""DOCX export helpers for generation Markdown output."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.generation_runs import GenerationRunRecord
from app.core.generation_template_completeness import GenerationTemplateCompletenessAssessment

GENERATION_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
GENERATION_DOCX_EXPORT_READY = "ready"
GENERATION_DOCX_EXPORT_WARNING = "warning"
GENERATION_DOCX_EXPORT_REASON_GENERATION_TRUNCATED = "generation_truncated"
GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_WARNING = "answer_quality_warning"
GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_FAILED = "answer_quality_failed"
GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_WARNING = "template_completeness_warning"
GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_FAILED = "template_completeness_failed"

_MARKDOWN_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MARKDOWN_ORDERED_LIST_PATTERN = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")
_MARKDOWN_UNORDERED_LIST_PATTERN = re.compile(r"^\s*[-*]\s+(.+?)\s*$")
_MARKDOWN_TABLE_SEPARATOR_PATTERN = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")

GENERATION_DOCX_STYLE_DEFAULT = "default"


@dataclass(frozen=True)
class GenerationDocxStyleProfile:
    document_type: str
    font_name: str
    heading_color: RGBColor
    table_style: str
    category: str


@dataclass(frozen=True)
class GenerationDocxExportReadiness:
    status: str
    reason_codes: tuple[str, ...]
    generation_truncated: bool
    answer_quality_status: str
    template_completeness_status: str


_DEFAULT_STYLE_PROFILE = GenerationDocxStyleProfile(
    document_type=GENERATION_DOCX_STYLE_DEFAULT,
    font_name="Malgun Gothic",
    heading_color=RGBColor(31, 41, 55),
    table_style="Table Grid",
    category="NeX-PCX Generation",
)
_STYLE_PROFILES = {
    "report": GenerationDocxStyleProfile(
        document_type="report",
        font_name="Malgun Gothic",
        heading_color=RGBColor(30, 64, 175),
        table_style="Light Shading Accent 1",
        category="NeX-PCX Report",
    ),
    "proposal": GenerationDocxStyleProfile(
        document_type="proposal",
        font_name="Malgun Gothic",
        heading_color=RGBColor(4, 120, 87),
        table_style="Light Shading Accent 4",
        category="NeX-PCX Proposal",
    ),
    "summary": GenerationDocxStyleProfile(
        document_type="summary",
        font_name="Malgun Gothic",
        heading_color=RGBColor(91, 33, 182),
        table_style="Light Shading Accent 5",
        category="NeX-PCX Summary",
    ),
    "meeting_minutes": GenerationDocxStyleProfile(
        document_type="meeting_minutes",
        font_name="Malgun Gothic",
        heading_color=RGBColor(180, 83, 9),
        table_style="Light Shading Accent 6",
        category="NeX-PCX Meeting Minutes",
    ),
}


def markdown_to_docx_bytes(
    markdown: str,
    *,
    title: str | None = None,
    document_type: str | None = None,
) -> bytes:
    """Convert a practical subset of Markdown to a DOCX byte payload."""

    document = Document()
    style_profile = generation_docx_style_profile(document_type)
    _configure_document(document, style_profile)
    if title:
        document.core_properties.title = title
    document.core_properties.subject = style_profile.document_type
    document.core_properties.category = style_profile.category
    _add_markdown(document, markdown, style_profile)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def generation_docx_style_profile(document_type: str | None) -> GenerationDocxStyleProfile:
    normalized_document_type = (document_type or "").strip().lower()
    return _STYLE_PROFILES.get(normalized_document_type, _DEFAULT_STYLE_PROFILE)


def assess_generation_docx_export_readiness(
    run: GenerationRunRecord,
    template_completeness: GenerationTemplateCompletenessAssessment,
) -> GenerationDocxExportReadiness:
    reason_codes: list[str] = []
    generation_truncated = _generation_run_truncated(run)
    answer_quality_status = _generation_run_answer_quality_status(run)
    template_completeness_status = template_completeness.status

    if generation_truncated:
        reason_codes.append(GENERATION_DOCX_EXPORT_REASON_GENERATION_TRUNCATED)
    if answer_quality_status == "failed":
        reason_codes.append(GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_FAILED)
    elif answer_quality_status == "warning":
        reason_codes.append(GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_WARNING)
    if template_completeness_status == "failed":
        reason_codes.append(GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_FAILED)
    elif template_completeness_status == "warning":
        reason_codes.append(GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_WARNING)

    return GenerationDocxExportReadiness(
        status=(GENERATION_DOCX_EXPORT_WARNING if reason_codes else GENERATION_DOCX_EXPORT_READY),
        reason_codes=tuple(reason_codes),
        generation_truncated=generation_truncated,
        answer_quality_status=answer_quality_status,
        template_completeness_status=template_completeness_status,
    )


def generation_docx_export_readiness_payload(
    readiness: GenerationDocxExportReadiness,
) -> dict[str, object]:
    return {
        "status": readiness.status,
        "reason_codes": list(readiness.reason_codes),
        "generation_truncated": readiness.generation_truncated,
        "answer_quality_status": readiness.answer_quality_status,
        "template_completeness_status": readiness.template_completeness_status,
    }


def _generation_run_truncated(run: GenerationRunRecord) -> bool:
    truncation = run.response_metadata.get("truncation")
    if isinstance(truncation, dict) and truncation.get("truncated") is True:
        return True
    return run.finish_reason == "length"


def _generation_run_answer_quality_status(run: GenerationRunRecord) -> str:
    answer_quality = run.response_metadata.get("answer_quality")
    if isinstance(answer_quality, dict):
        status = answer_quality.get("status")
        if isinstance(status, str) and status.strip():
            return status.strip()
    guardrail_status = run.guardrail_metadata.get("answer_quality_status")
    if isinstance(guardrail_status, str) and guardrail_status.strip():
        return guardrail_status.strip()
    return "not_available"


def _configure_document(
    document: DocxDocument,
    style_profile: GenerationDocxStyleProfile,
) -> None:
    for section in document.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    _set_style_font(document, "Normal", style_profile.font_name, 10.5)
    _set_style_font(document, "Heading 1", style_profile.font_name, 18, style_profile.heading_color)
    _set_style_font(document, "Heading 2", style_profile.font_name, 14, style_profile.heading_color)
    _set_style_font(document, "Heading 3", style_profile.font_name, 12, style_profile.heading_color)
    _set_style_font(document, "Heading 4", style_profile.font_name, 11, style_profile.heading_color)


def _set_style_font(
    document: DocxDocument,
    style_name: str,
    font_name: str,
    size_pt: float,
    color: RGBColor | None = None,
) -> None:
    style = document.styles[style_name]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = color
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_markdown(
    document: DocxDocument,
    markdown: str,
    style_profile: GenerationDocxStyleProfile,
) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code_block = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                _add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue
        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if _is_markdown_table_start(lines, index):
            table_lines: list[str] = []
            while index < len(lines) and _is_markdown_table_line(lines[index]):
                table_lines.append(lines[index])
                index += 1
            _add_table(document, table_lines, style_profile)
            continue

        heading_match = _MARKDOWN_HEADING_PATTERN.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(_plain_text(heading_match.group(2)), level=level)
            index += 1
            continue

        unordered_match = _MARKDOWN_UNORDERED_LIST_PATTERN.match(line)
        if unordered_match:
            document.add_paragraph(_plain_text(unordered_match.group(1)), style="List Bullet")
            index += 1
            continue

        ordered_match = _MARKDOWN_ORDERED_LIST_PATTERN.match(line)
        if ordered_match:
            document.add_paragraph(_plain_text(ordered_match.group(1)), style="List Number")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and _is_paragraph_continuation(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        document.add_paragraph(_plain_text(" ".join(paragraph_lines)))

    if code_lines:
        _add_code_block(document, code_lines)


def _is_paragraph_continuation(line: str) -> bool:
    stripped = line.strip()
    return bool(
        stripped
        and not stripped.startswith("```")
        and not _MARKDOWN_HEADING_PATTERN.match(stripped)
        and not _MARKDOWN_UNORDERED_LIST_PATTERN.match(line)
        and not _MARKDOWN_ORDERED_LIST_PATTERN.match(line)
        and not _is_markdown_table_line(line)
    )


def _is_markdown_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and _is_markdown_table_line(lines[index])
        and _MARKDOWN_TABLE_SEPARATOR_PATTERN.match(lines[index + 1].strip()) is not None
    )


def _is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _add_table(
    document: DocxDocument,
    table_lines: list[str],
    style_profile: GenerationDocxStyleProfile,
) -> None:
    rows = [
        _split_table_row(line)
        for line in table_lines
        if _MARKDOWN_TABLE_SEPARATOR_PATTERN.match(line.strip()) is None
    ]
    rows = [row for row in rows if row]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    try:
        table.style = style_profile.table_style
    except KeyError:
        table.style = _DEFAULT_STYLE_PROFILE.table_style
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = _plain_text(row[column_index]) if column_index < len(row) else ""


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_code_block(document: DocxDocument, code_lines: list[str]) -> None:
    if not code_lines:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Courier New"


def _plain_text(value: str) -> str:
    text = value.replace("**", "").replace("__", "").replace("*", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()
