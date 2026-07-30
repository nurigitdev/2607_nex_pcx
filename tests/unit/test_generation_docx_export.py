from datetime import UTC, datetime
from io import BytesIO

import pytest
from docx import Document
from docx.shared import RGBColor

import app.core.generation_docx_export as generation_docx_export
from app.core.generation_docx_export import (
    GENERATION_DOCX_EXPORT_READY,
    GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_FAILED,
    GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_WARNING,
    GENERATION_DOCX_EXPORT_REASON_GENERATION_TRUNCATED,
    GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_FAILED,
    GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_WARNING,
    GENERATION_DOCX_EXPORT_WARNING,
    GENERATION_DOCX_STYLE_DEFAULT,
    GenerationDocxStyleProfile,
    assess_generation_docx_export_readiness,
    generation_docx_export_evidence_from_run,
    generation_docx_export_evidence_payload,
    generation_docx_export_readiness_payload,
    generation_docx_style_profile,
    markdown_to_docx_bytes,
)
from app.core.generation_runs import GenerationRunRecord
from app.core.generation_template_completeness import GenerationTemplateCompletenessAssessment

NOW = datetime(2026, 7, 27, 12, 30, tzinfo=UTC)


def test_markdown_to_docx_bytes_converts_headings_lists_tables_and_code() -> None:
    docx_bytes = markdown_to_docx_bytes(
        "\n".join(
            (
                "# 보고서 초안",
                "",
                "## 요약",
                "첫 문단입니다.",
                "",
                "- 항목 A",
                "1. 항목 B",
                "",
                "| 구분 | 값 |",
                "| --- | --- |",
                "| 상태 | 정상 |",
                "",
                "```json",
                '{"ok": true}',
                "```",
            )
        ),
        title="Pytest DOCX",
        document_type="report",
    )

    document = Document(BytesIO(docx_bytes))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    table_text = [
        [cell.text for cell in row.cells] for table in document.tables for row in table.rows
    ]

    assert document.core_properties.title == "Pytest DOCX"
    assert document.core_properties.subject == "report"
    assert document.core_properties.category == "NeX-PCX Report"
    assert paragraph_text[:5] == ["보고서 초안", "요약", "첫 문단입니다.", "항목 A", "항목 B"]
    assert '{"ok": true}' in paragraph_text[-1]
    assert table_text == [["구분", "값"], ["상태", "정상"]]


def test_markdown_to_docx_bytes_ignores_empty_tables_and_unclosed_code_blocks() -> None:
    docx_bytes = markdown_to_docx_bytes(
        "\n".join(
            (
                "plain paragraph",
                "",
                "| --- | --- |",
                "",
                "```",
                "unclosed code",
            )
        )
    )

    document = Document(BytesIO(docx_bytes))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    assert "plain paragraph" in paragraph_text
    assert "unclosed code" in paragraph_text
    assert document.tables == []


def test_markdown_to_docx_bytes_handles_regression_markdown_boundaries() -> None:
    docx_bytes = markdown_to_docx_bytes(
        "\n".join(
            (
                "# 회귀 fixture",
                "첫 번째 줄입니다.",
                "이어지는 줄입니다.",
                "",
                "```",
                "```",
                "",
                "### 다음 섹션",
                "문단입니다.",
            )
        )
    )

    document = Document(BytesIO(docx_bytes))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    assert "첫 번째 줄입니다. 이어지는 줄입니다." in paragraph_text
    assert "" not in paragraph_text
    assert paragraph_text[-2:] == ["다음 섹션", "문단입니다."]


def test_generation_docx_style_profile_falls_back_for_unknown_document_type() -> None:
    profile = generation_docx_style_profile(" unknown ")

    assert profile.document_type == GENERATION_DOCX_STYLE_DEFAULT
    assert profile.category == "NeX-PCX Generation"


def test_generation_docx_export_readiness_passes_for_clean_run() -> None:
    readiness = assess_generation_docx_export_readiness(_run(), _template_completeness())
    payload = generation_docx_export_readiness_payload(readiness)

    assert readiness.status == GENERATION_DOCX_EXPORT_READY
    assert readiness.reason_codes == ()
    assert payload == {
        "status": "ready",
        "reason_codes": [],
        "generation_truncated": False,
        "answer_quality_status": "passed",
        "template_completeness_status": "passed",
    }


def test_generation_docx_export_readiness_warns_for_truncation_and_quality_failure() -> None:
    readiness = assess_generation_docx_export_readiness(
        _run(
            finish_reason="length",
            response_metadata={"answer_quality": {"status": "failed"}},
        ),
        _template_completeness(status="failed"),
    )

    assert readiness.status == GENERATION_DOCX_EXPORT_WARNING
    assert readiness.generation_truncated is True
    assert readiness.reason_codes == (
        GENERATION_DOCX_EXPORT_REASON_GENERATION_TRUNCATED,
        GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_FAILED,
        GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_FAILED,
    )


def test_generation_docx_export_readiness_warns_from_metadata_fallbacks() -> None:
    readiness = assess_generation_docx_export_readiness(
        _run(
            response_metadata={"truncation": {"truncated": True}},
            guardrail_metadata={"answer_quality_status": "warning"},
        ),
        _template_completeness(status="warning"),
    )

    assert readiness.status == GENERATION_DOCX_EXPORT_WARNING
    assert readiness.answer_quality_status == "warning"
    assert readiness.reason_codes == (
        GENERATION_DOCX_EXPORT_REASON_GENERATION_TRUNCATED,
        GENERATION_DOCX_EXPORT_REASON_ANSWER_QUALITY_WARNING,
        GENERATION_DOCX_EXPORT_REASON_TEMPLATE_COMPLETENESS_WARNING,
    )


def test_generation_docx_export_readiness_marks_missing_quality_as_not_available() -> None:
    readiness = assess_generation_docx_export_readiness(
        _run(response_metadata={}), _template_completeness()
    )

    assert readiness.status == GENERATION_DOCX_EXPORT_READY
    assert readiness.answer_quality_status == "not_available"


def test_generation_docx_export_readiness_falls_back_when_answer_quality_status_is_blank() -> None:
    readiness = assess_generation_docx_export_readiness(
        _run(
            response_metadata={"answer_quality": {"status": "  "}},
            guardrail_metadata={"answer_quality_status": "passed"},
        ),
        _template_completeness(),
    )

    assert readiness.status == GENERATION_DOCX_EXPORT_READY
    assert readiness.answer_quality_status == "passed"


def test_markdown_to_docx_bytes_embeds_export_evidence_metadata() -> None:
    run = _run()
    readiness = assess_generation_docx_export_readiness(run, _template_completeness())
    evidence = generation_docx_export_evidence_from_run(
        run,
        template={
            "template_key": "report",
            "template_version": "v3",
            "document_type": "report",
        },
        readiness=readiness,
    )
    docx_bytes = markdown_to_docx_bytes(
        "# 보고서",
        document_type="report",
        export_evidence=evidence,
    )

    document = Document(BytesIO(docx_bytes))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    evidence_table = {row.cells[0].text: row.cells[1].text for row in document.tables[-1].rows}
    payload = generation_docx_export_evidence_payload(evidence)

    assert document.core_properties.author == "NeX-PCX"
    assert "generation_run_id=42" in document.core_properties.keywords
    assert "readiness=ready" in document.core_properties.comments
    assert paragraph_text[-1] == "Export Evidence"
    assert evidence_table["Generation Run ID"] == "42"
    assert evidence_table["Search Log ID"] == "24"
    assert evidence_table["Provider"] == "mock_qwen35_122b_a10b_nvfp4 (mock)"
    assert evidence_table["Template"] == "report / v3"
    assert evidence_table["Export Readiness"] == "ready"
    assert evidence_table["Readiness Reasons"] == "-"
    assert payload["generation_run_id"] == 42
    assert payload["export_readiness_reasons"] == []


def test_markdown_to_docx_bytes_falls_back_when_profile_table_style_is_unavailable(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    broken_profile = GenerationDocxStyleProfile(
        document_type="broken",
        font_name="Malgun Gothic",
        heading_color=RGBColor(31, 41, 55),
        table_style="Definitely Missing Table Style",
        category="Broken",
    )
    monkeypatch.setitem(generation_docx_export._STYLE_PROFILES, "broken", broken_profile)

    docx_bytes = markdown_to_docx_bytes(
        "\n".join(
            (
                "| 구분 | 값 |",
                "| --- | --- |",
                "| 상태 | 정상 |",
            )
        ),
        document_type="broken",
        export_evidence=generation_docx_export_evidence_from_run(
            _run(),
            template={"template_key": "broken", "template_version": "", "document_type": ""},
            readiness=assess_generation_docx_export_readiness(_run(), _template_completeness()),
        ),
    )

    document = Document(BytesIO(docx_bytes))

    assert len(document.tables) == 2
    assert document.tables[0].rows[1].cells[1].text == "정상"
    assert document.tables[1].rows[-1].cells[1].text == "-"


def _run(
    *,
    finish_reason: str = "stop",
    response_metadata: dict[str, object] | None = None,
    guardrail_metadata: dict[str, object] | None = None,
) -> GenerationRunRecord:
    return GenerationRunRecord(
        generation_run_id=42,
        search_log_id=24,
        retrieval_package_key="pkg-generation-42",
        generation_template_id=None,
        provider_config_id=7,
        provider_name="mock_qwen35_122b_a10b_nvfp4",
        provider_mode="mock",
        model_id="nvidia/Qwen3.5-122B-A10B-NVFP4",
        prompt_version="grounded_answer_v1",
        prompt_hash="prompt-hash",
        context_hash="context-hash",
        status="succeeded",
        guardrail_status="allowed",
        retrieval_confidence_status="answerable",
        citation_readiness_status="ready",
        query_text="내부 규정 요약",
        answer_text="요약 답변입니다. [RCP-001]",
        finish_reason=finish_reason,
        input_token_count=100,
        output_token_count=20,
        total_token_count=120,
        elapsed_ms=12,
        request_metadata={},
        response_metadata=(
            response_metadata
            if response_metadata is not None
            else {"answer_quality": {"status": "passed"}}
        ),
        guardrail_metadata=guardrail_metadata or {},
        error_message=None,
        created_by="pytest",
        created_by_user_id=1,
        started_at=NOW,
        finished_at=NOW,
        created_at=NOW,
        updated_at=NOW,
    )


def _template_completeness(status: str = "passed") -> GenerationTemplateCompletenessAssessment:
    return GenerationTemplateCompletenessAssessment(
        status=status,
        template_key="grounded_answer",
        template_name="근거 기반 답변",
        template_version="v1",
        document_type="answer",
        output_format="markdown",
        required_section_count=1,
        present_required_section_count=1 if status != "failed" else 0,
        citation_required_section_count=1,
        cited_required_section_count=1 if status != "failed" else 0,
        required_section_coverage_percent=100.0 if status != "failed" else 0.0,
        required_section_citation_coverage_percent=100.0 if status != "failed" else 0.0,
        section_checks=(),
        reason_codes=(),
    )
